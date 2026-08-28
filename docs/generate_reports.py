import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Helpers
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=130, right=130):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.runs[0]
    if level == 1:
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x4B, 0x87) # Cisco Navy
    elif level == 2:
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    elif level == 3:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return h

def add_callout(doc, text, title=None, border_color="004B87", bg_color="F0F4F8"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if title:
        r_title = p.add_run(f"{title}\n")
        r_title.bold = True
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(10.5)
        r_title.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:color="CBD5E1"/><w:left w:val="single" w:sz="4" w:color="CBD5E1"/><w:bottom w:val="single" w:sz="4" w:color="CBD5E1"/><w:right w:val="single" w:sz="4" w:color="CBD5E1"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

def add_table_data(doc, headers, rows_data, col_widths=None):
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    # Header Row
    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "004B87")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    # Data Rows
    for r_idx, row in enumerate(rows_data):
        row_cells = tbl.rows[r_idx + 1].cells
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                
    # Borders
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:color="E2E8F0"/><w:left w:val="single" w:sz="4" w:color="E2E8F0"/><w:bottom w:val="single" w:sz="4" w:color="E2E8F0"/><w:right w:val="single" w:sz="4" w:color="E2E8F0"/></w:tcBorders>')
            tcPr.append(borders)
            
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
                
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

def add_p(doc, text, bold_prefix=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.bold = True
        r_bold.font.name = 'Calibri'
        r_bold.font.size = Pt(10)
        r_bold.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.bold = True
        r_bold.font.name = 'Calibri'
        r_bold.font.size = Pt(9.5)
        r_bold.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

def create_base_doc():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
    return doc

def add_doc_header(doc, title, subtitle, meta_dict):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(8)
    run_sub = p_sub.add_run(subtitle)
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    meta_rows = [[k, v] for k, v in meta_dict.items()]
    add_table_data(doc, ["Project Metadata Field", "Specification Details"], meta_rows, col_widths=[2.3, 4.2])

print("Helper definitions loaded successfully.")
